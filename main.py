import os
from docx import Document

def merge_docx_files(directory):
    merged_document = Document()

    sorted_docx_files = [f for f in sorted(os.listdir(directory)) if f.endswith('.docx')]

    for filename in sorted_docx_files:
        file_path = os.path.join(directory, filename)

        current_document = Document(file_path)

        if len(merged_document.paragraphs) > 0:
            merged_document.add_page_break()

        for para in current_document.paragraphs:
            new_para = merged_document.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                # Font attributes
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.color.rgb = run.font.color.rgb
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
            new_para.style = para.style

    # Save
    merged_document.save(os.path.join(directory, 'merged_document.docx'))
    print('Successfully merged into merged_document.docx')

directory = 'docx_files'
merge_docx_files(directory)